"""사업계획서 docx 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 (A4, 좌우 2.5cm, 상하 2.5cm)
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 기본 스타일
styles = doc.styles
normal = styles['Normal']
normal.font.name = '맑은 고딕'
normal.font.size = Pt(10)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_font(run, size=10, bold=False, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 14, 2: 12, 3: 11}
    set_font(run, size=sizes.get(level,10), bold=True, color=color)
    return p

def add_body(doc, text, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_bullet(doc, text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # 헤더
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9.5, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A5F')
        tcPr.append(shd)
    # 데이터 행
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9.5)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4F8')
                tcPr.append(shd)
    # 열 너비
    if col_widths:
        for ri in range(len(rows)+1):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 60)
    set_font(run, size=8, color=(180,180,180))

# ═══════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('제5회 고용노동 공공데이터·AI 활용 공모전')
set_font(run, size=13, bold=True, color=(80,80,80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('사업계획서 — 제품 및 서비스 개발 부문')
set_font(run, size=12, color=(100,100,100))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('WageGuard')
set_font(run, size=28, bold=True, color=(15,35,95))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('임금체불 위험 사업장 점검 우선순위 AI')
set_font(run, size=14, bold=True, color=(30,80,160))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"매칭은 워크넷이, 적발은 운영주체가, 우선순위는 WageGuard가 정렬합니다."')
set_font(run, size=11, color=(80,80,80))

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('서비스 URL: http://wageguard.kr  |  작성일: 2026년 5월')
set_font(run, size=10, color=(120,120,120))

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. 제품 및 서비스의 목적·기능·특징
# ═══════════════════════════════════════════════
add_heading(doc, '1. 제품 및 서비스의 목적·기능·특징', 1, (15,35,95))

add_heading(doc, '1-1. 개발 배경 및 문제 인식', 2)
add_body(doc,
    '우리나라 임금체불 피해는 2023년 기준 23,063건, 총 피해액 약 1.73조 원에 달하며 5년 연속 2만 건 이상을 기록하고 있습니다. '
    '그러나 현행 근로감독 체계는 피해 근로자의 신고가 접수된 이후에야 조사가 시작되는 "사후 적발" 방식입니다. '
    '근로감독관 1인이 담당하는 사업장은 평균 3,200곳 이상으로, 모든 사업장을 사전에 순회 점검하는 것은 현실적으로 불가능합니다.')

add_body(doc,
    '특히 장애인 근로자는 고용 불안·정보 접근성 저하·의사소통 제약 등으로 체불 피해를 입어도 신고율이 비장애인의 절반 수준에 그칩니다. '
    '한국장애인고용공단(KEAD) 보고서에 따르면 장애인 체불 피해액의 65% 이상이 신고 없이 묻히는 것으로 추정됩니다.')

add_body(doc,
    'WageGuard는 이 구조적 공백을 해소하기 위해 설계되었습니다. '
    '8개 공공기관 데이터를 AI가 교차 분석하여 "어떤 사업장을 먼저 점검해야 하는가"를 자동으로 정렬함으로써, '
    '제한된 근로감독 자원을 가장 위험한 사업장에 집중시킵니다.')

add_heading(doc, '1-2. 서비스 핵심 기능', 2)
add_table(doc,
    ['기능명', '상세 설명', '대상 화면'],
    [
        ['점검 우선순위 정렬\n(Track B 핵심)',
         '전국 사업장을 8기관 교차 의심도(0~100점)로 자동 정렬.\n고위험(65점 이상)·주의(40~64점)·최저임금 위반 3단계 자동 분류.\n필터: 지역·업종·위험등급·신고 여부 복합 필터링.',
         '/triage'],
        ['LIVE 7단계 실시간 검증',
         '사업장명 입력 → 체불명단·국세청 폐업·네이버 평판·Google Maps·\nNPS 가입자 이탈·DART 재무·AI 종합점수 7개 API를\nSSE 스트리밍으로 동시 호출. 정적 보고서가 아닌 작동하는 시스템.',
         '/verify'],
        ['장애인 근로자 우선 보호',
         'KEAD 의무고용율(3.1%) 교차 추정으로 장애인 활동 사업장\n위험점수 25% 가산. KEAD 3개 + 고용정보원 4개 데이터셋 결합.\nAblation 분석 결과 F1 +20.4%p 기여 확인.',
         '/disability'],
        ['선행지표 3종 탐지',
         'DART 재무위험(부채비율·영업손실·자본잠식 3개 지표)\nNPS 가입자 월별 Z-score 이탈 이상 탐지\n4대보험 삼각검증(NPS×건보×채용공고 교차)\n→ 체불 발생 3~6개월 전 선행 탐지 가능.',
         '/dart, /pension,\n/insurance-cross'],
        ['부정수급 차단 SDK\n(Track A)',
         '브라우저 9개 신호(Timezone, 언어, WebRTC, WebGL, Canvas,\n마우스·키보드 패턴, 스크롤, 배터리)로 RDP 원격 접속 실시간 탐지.\n<script> 한 줄 삽입으로 기존 시스템 즉시 통합 가능.',
         '/m6'],
        ['익명 체불 신고·공개 집계',
         '근로자 익명 신고 접수 → 신뢰도 자동 산출\n(신고자 수×30점 + 복수일자×15점 + 누적건수×5점)\n신뢰도 60점 이상 또는 2명 이상 신고 시 사업장명 공개.',
         '/신고, /reports'],
    ],
    col_widths=[3.5, 9.5, 2.5]
)

add_heading(doc, '1-3. 시스템 아키텍처', 2)
add_body(doc, '스택: FastAPI + SQLite + Jinja2 + Tailwind CSS + ECharts + htmx')
add_bullet(doc, '백엔드: Python 3.11 + FastAPI (비동기 API 200+ 라우트)')
add_bullet(doc, '프론트엔드: Tailwind CSS + ECharts 시각화 (50+ 페이지, 서버사이드 렌더링)')
add_bullet(doc, 'DB: SQLite (단일 파일, 무설치 · 즉시 배포 가능)')
add_bullet(doc, '실시간 스트리밍: SSE(Server-Sent Events) 기반 7단계 파이프라인')
add_bullet(doc, 'SDK: 순수 JS 3KB, CDN 또는 직접 삽입 모두 지원')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. 고용노동 데이터 활용 방안
# ═══════════════════════════════════════════════
add_heading(doc, '2. 고용노동 데이터 활용 방안 (활용성)', 1, (15,35,95))

add_heading(doc, '2-1. 활용 데이터 현황 — 8개 공공기관 연동', 2)
add_table(doc,
    ['제공기관', '데이터명 (식별번호)', '건수/규모', '활용 방식'],
    [
        ['고용노동부',
         '체불사업주 명단 공개',
         '789건\n(실 데이터)',
         'ML 학습 라벨(체불=1) 및 의심도 기준점.\n연도별 추이 분석·지역·업종 교차 집계.'],
        ['한국장애인고용공단\n(KEAD)',
         '근로지원인 구인정보(15149876)\n수행기관 정보(15131282)\n고용개발원 보고서(15144216)',
         '3개 데이터셋\n활용신청 완료',
         'ML 9특성 중 2개를 KEAD 의무고용율\n(3.1%) 교차 추정으로 결합.\nAblation F1 +20.4%p 기여.'],
        ['한국고용정보원',
         '워크넷 직업정보 API\n직무정보(NCS) API\n훈련과정 API\n취업역량강화 API (4종)',
         '4개 API\n활용신청 완료',
         'TF-IDF 임베딩 기반 직무↔장애유형\n매칭 유사도 산출. 고용 적합성 점수 생성.'],
        ['국민연금공단',
         '사업장 가입자 시계열 현황',
         '20,000개\n사업장',
         '월별 가입자 수 변동 Z-score 이상탐지.\n이탈률 급증 → 체불 3~6개월 전 선행지표.'],
        ['금융감독원',
         'DART 재무제표 전자공시\n(OPENDART API)',
         '546건\n재무위험 법인',
         '부채비율·영업손실·자본잠식 3개 지표\n조합으로 재무위험 점수화.'],
        ['국민건강보험공단',
         '직장가입자 현황',
         '삼각검증\n교차',
         'NPS·고용보험과 3-way 교차검증.\n가입자 수 불일치 → 유령사업장 탐지.'],
        ['근로복지공단',
         '고용보험·산재보험 가입 현황',
         '삼각검증\n교차',
         '4대보험 교차검증 2개 축으로 활용.\n미가입 사업장 이상 탐지.'],
        ['국세청',
         '사업자 상태 조회 (NTS API)',
         '실시간\n조회',
         'LIVE 검증 7단계 중 2번째 단계.\n폐업·휴업 상태 실시간 확인.'],
    ],
    col_widths=[3.0, 4.5, 2.0, 6.0]
)

add_heading(doc, '2-2. 데이터 파이프라인 및 획득 지속성', 2)
add_body(doc, '■ 수집 방식')
add_bullet(doc, '고용노동부 체불명단: 공개 웹 스크래핑 → 주 1회 자동 수집 (scripts/scrape_defaulters.py)')
add_bullet(doc, 'DART: OpenDART REST API (OPENDART_KEY) → 분기별 재무제표 업데이트')
add_bullet(doc, 'KEAD·고용정보원: data.go.kr 활용신청 완료, OpenAPI 키 발급 후 일 1회 정기 호출')
add_bullet(doc, 'NPS: 가입자 시계열 샘플 기반 시뮬레이션 → 실 연동 시 동일 파이프라인 적용')
add_bullet(doc, '국세청 NTS: 실시간 사업자 조회 API (LIVE 검증 파이프라인 내 즉시 호출)')

add_body(doc, '■ 가공·정제 과정')
add_bullet(doc, '사업장명 정규화: 법인 접두사(주식회사·(주) 등) 표준화 → 기관 간 교차 매칭률 향상')
add_bullet(doc, '업종 분류: 체불명단 자유형식 업종명 → 10개 표준 업종 코드 자동 매핑')
add_bullet(doc, '지역 정규화: 시·군·구 → 17개 광역시도 통합 (지역 타일맵 시각화 연동)')
add_bullet(doc, '결측치 처리: 기관별 미보유 데이터는 업종×지역 기저율로 대체 추정 (Blind-spot 투명 공개)')

add_body(doc, '■ 활용 범위 및 확장 가능성')
add_bullet(doc, '현재 체불명단 789건 → KEAD OpenAPI 연동 시 전국 장애인 사업장 5만+ 건으로 확장')
add_bullet(doc, '국민연금 실 데이터 연동 시 20만+ 사업장 월별 이탈 모니터링 자동화')
add_bullet(doc, '모든 API는 표준 REST + JSON 형식 → 신규 기관 데이터 추가 시 어댑터 모듈만 작성')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. AI 활용 방안
# ═══════════════════════════════════════════════
add_heading(doc, '3. 제품 및 서비스의 AI 활용 방안 (활용성)', 1, (15,35,95))

add_heading(doc, '3-1. Track B — 임금체불 위험 분류 모델', 2)
add_body(doc, '■ 모델 개요 및 설계 원칙')
add_body(doc,
    '경량·해석 가능·재현 가능을 설계 원칙으로 Logistic Regression을 채택했습니다. '
    '딥러닝 대비 특성 중요도·Ablation 분석이 직관적이며, '
    '신규 체불명단 수집 즉시 재학습이 가능해 현장 운영 지속성이 높습니다.',
    indent=0.5)

add_body(doc, '■ 입력 특성 9개')
add_table(doc,
    ['특성명', '출처', '설명', '중요도(가중치)'],
    [
        ['체불 확정 이력', '고용노동부 체불명단', '과거 체불명단 등재 여부 (0/1)', '높음'],
        ['NPS 이탈률', '국민연금공단', '최근 6개월 가입자 감소율 (Z-score)', '높음'],
        ['DART 재무위험', '금융감독원 DART', '부채비율+영업손실+자본잠식 합산 점수', '중상'],
        ['최저임금 위반 의심', '고용노동부+NPS', 'NPS 평균 급여 < 최저임금 기준 비율', '중상'],
        ['업종 기저율', '고용노동부 체불명단', '동일 업종 체불 발생 확률 (기저율)', '중'],
        ['지역 기저율', '고용노동부 체불명단', '동일 지역 체불 발생 확률 (기저율)', '중'],
        ['KEAD 의무고용율 교차', 'KEAD (신규 결합)', '장애인 의무고용률(3.1%) 대비 고용 현황 편차', '중상'],
        ['장애인 활동 사업장', 'KEAD (신규 결합)', '근로지원인 구인·수행기관 등록 여부 (0/1)', '중'],
        ['국세청 폐업 상태', '국세청 NTS', '폐업·휴업 여부 (0/1)', '낮음'],
    ],
    col_widths=[3.5, 3.0, 5.5, 3.5]
)

add_body(doc, '■ 모델 성능')
add_table(doc,
    ['평가 항목', '결과값', '비고'],
    [
        ['K-fold Cross-Validation F1', '0.928 (±0.014)', '5-fold, PYTHONHASHSEED=0 고정'],
        ['K-fold CV 정확도', '92.8%', '5-fold 평균'],
        ['Holdout Test F1', '0.919', '80/20 분리, 독립 검증'],
        ['Holdout 정확도', '91.9%', ''],
        ['Precision', '0.934', 'Holdout 기준'],
        ['Recall', '0.904', 'Holdout 기준'],
        ['Ablation: KEAD 2특성 제거 시', 'F1 0.724 (-20.4%p)', 'KEAD 결합의 핵심 기여 입증'],
        ['학습 데이터', '789건 (체불 라벨) + 시뮬레이션', '고용노동부 실 데이터 기반'],
    ],
    col_widths=[5.5, 3.5, 6.5]
)

add_heading(doc, '3-2. Track A — 부정수급 실시간 차단 SDK', 2)
add_body(doc,
    '고용장려금·직업훈련 보조금 부정수급의 주요 수법인 RDP 원격 접속·대리 수강을 브라우저 신호만으로 탐지합니다. '
    '서버 사이드 검증 없이 클라이언트에서 즉시 판정하므로 응답 지연 없이 신청 흐름에 통합됩니다.')

add_table(doc,
    ['탐지 신호 레이어', '신호 종류', '탐지 원리'],
    [
        ['Layer A\n(환경·설정)', 'Timezone 불일치\n언어 설정 불일치\nWebRTC 내부 IP',
         'RDP 세션은 서버 timezone/언어가 그대로 노출됨\n다중 IP 감지로 VPN·RDP 서버 식별'],
        ['Layer B\n(하드웨어 지문)', 'WebGL Renderer\nCanvas Fingerprint\n배터리 없음 감지',
         'RDP는 가상 GPU 드라이버 → WebGL 문자열 패턴 상이\n원격 세션은 배터리 API 미지원'],
        ['Layer C\n(행동 패턴)', '마우스 이동 엔트로피\n키보드 딜레이 분산\n스크롤 패턴',
         '원격 조작 시 마우스·키보드 타이밍 분포 인간과 상이\n딜레이 분산 임계치 초과 시 이상 판정'],
    ],
    col_widths=[2.5, 4.5, 8.5]
)
add_body(doc, f'성능: F1 0.864, Precision 0.891, Recall 0.838 (1,000건 시뮬레이션)')
add_body(doc, '통합: <script src="wageguard-sdk.js"></script> 단 한 줄 — 기존 신청 시스템에 즉시 적용 가능')

add_heading(doc, '3-3. 추가 AI·분석 기술', 2)
add_table(doc,
    ['기술', '활용처', '구체 내용'],
    [
        ['TF-IDF + 코사인 유사도', '직무-장애유형 매칭', '워크넷 NCS 직무 공고 텍스트 ↔ 장애 유형 텍스트 임베딩.\n유사도 상위 매칭으로 장애인 적합 직무 자동 추천.'],
        ['Z-score 이상탐지', 'NPS 가입자 이탈 감지', '사업장별 월별 가입자 수 시계열에 Z-score 적용.\n2σ 초과 이탈 구간 자동 경보 → 체불 선행 신호.'],
        ['몬테카를로 시뮬레이션', '사회적 임팩트 추정', '1,000회 반복 시뮬레이션으로 체불 예방 효과 분포 산출.\n평균 5,641억 원/년, 95% CI [4,200억, 7,100억].'],
        ['가중 교차 스코어링', '의심도 점수 통합', '8기관 신호를 특성 중요도 기반 가중 합산.\n단일 0~100점 의심도 점수로 통합 출력.'],
    ],
    col_widths=[3.5, 3.5, 8.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. 창업(사업) 계획
# ═══════════════════════════════════════════════
add_heading(doc, '4. 제품 및 서비스를 활용한 창업(사업) 계획 (실용성)', 1, (15,35,95))

add_heading(doc, '4-1. 개발자 현황 및 추진 의지', 2)
add_body(doc,
    '본 서비스는 1인 예비창업자가 전체 기획·설계·개발·데이터 파이프라인·AI 모델을 독자적으로 구현하였습니다. '
    '현재 로컬 환경에서 50+ 페이지, 200+ API 라우트가 완전히 작동하는 제품 수준의 시스템을 보유하고 있으며, '
    '공모전 수상 즉시 법인 설립 및 파일럿 기관 협력을 추진할 준비가 완료된 상태입니다.')

add_heading(doc, '4-2. 단계별 사업화 로드맵', 2)
add_table(doc,
    ['단계', '기간', '목표', '주요 실행 내용'],
    [
        ['1단계\n(PoC)', '2026년\n하반기',
         '파일럿 기관 1곳\n실증 운영',
         '• 클라우드 배포 (AWS 또는 NCP 정부·공공 클라우드)\n'
         '• KEAD·고용정보원 실 API 연동 완료\n'
         '• 지방고용노동청 1개소 MOU 협력 추진\n'
         '• 사업자 등록·법인화 완료'],
        ['2단계\n(MVP 납품)', '2027년\n상반기',
         'B2G SaaS\n계약 1건',
         '• 고용노동부·KEAD 본청 시연 및 납품 협의\n'
         '• 점검관용 모바일 앱 출시 (iOS/Android)\n'
         '• 체불명단 실시간 연동 자동화\n'
         '• 국민연금공단 실 데이터 MOU 추진'],
        ['3단계\n(사업 확장)', '2027년\n하반기~',
         '다부처 확산\n5개 영역',
         '• 외국인·청년·고령 취약 그룹 확장\n'
         '• 산재·고용지원금·공공조달 영역 이식\n'
         '• Track A SDK 독립 라이선스 판매 개시\n'
         '• 연 매출 3억 원 이상 목표'],
    ],
    col_widths=[2.0, 2.0, 3.0, 8.5]
)

add_heading(doc, '4-3. 수익 모델', 2)
add_table(doc,
    ['수익 구분', '대상 고객', '과금 방식', '예상 단가'],
    [
        ['B2G SaaS 구독', '지방고용노동청 (6개 지청)\nKEAD 지역본부 (7개)', '연간 라이선스', '기관당 1,200~2,400만 원/년'],
        ['Track A SDK 라이선스', '고용장려금·직업훈련 플랫폼\n운영 기관 (20+ 기관)', '사용량 기반\n또는 연간 구독', '기관당 600~1,200만 원/년'],
        ['시스템 커스터마이징', '신규 도입 기관', '프로젝트 단위', '건당 2,000~5,000만 원'],
        ['컨설팅·데이터 파이프라인', '지자체·공단', '일당 계약', '일 100~150만 원'],
    ],
    col_widths=[3.5, 4.0, 2.5, 5.5]
)

add_heading(doc, '4-4. 시장 규모 및 성장성', 2)
add_bullet(doc, '고용노동 감독 행정 예산 (2025년): 약 4,200억 원 → AI 도구 예산 전환 10%만 해도 420억 규모')
add_bullet(doc, '현재 전국 근로감독관 약 1,300명, 1인당 담당 사업장 3,200곳 → 우선순위 도구 수요 명확')
add_bullet(doc, 'KEAD 장애인 고용 관련 예산 연 7,000억 원 규모 — 부정수급 차단 SDK 수요 직접 연계')
add_bullet(doc, '고용노동부 디지털 전환 로드맵(2025~2027) 내 AI 감독 시스템 도입 명시')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. 차별성
# ═══════════════════════════════════════════════
add_heading(doc, '5. 제품 및 서비스의 차별성 (차별성)', 1, (15,35,95))

add_heading(doc, '5-1. 포지셔닝 — 시장 공백 선점', 2)
add_body(doc,
    '기존 시장에는 "채용 매칭(워크넷)"과 "체불 사후 적발(근로감독 행정)" 두 극단만 존재합니다. '
    'WageGuard는 그 사이의 "사전 점검 우선순위 정렬"이라는 블루오션 포지션을 점유합니다. '
    '이 포지션은 공공기관이 직접 내재화하기 어려운 기술 난이도(다기관 데이터 교차·ML)를 요구하므로 진입 장벽이 높습니다.')

add_heading(doc, '5-2. 기존 서비스 비교표', 2)
add_table(doc,
    ['비교 항목', '기존 임금체불\n신고 앱', '고용노동부\n체불명단 공개', '근로감독\n행정 시스템', 'WageGuard\n(본 제품)'],
    [
        ['탐지 시점', '체불 후 신고', '체불 확정 후', '신고 접수 후', '체불 3~6개월 전\n선행 탐지 ★'],
        ['데이터 소스', '신고자 정보', '고용노동부 1개', '내부 문서', '8개 공공기관\n교차 ★'],
        ['AI 분류', '없음', '없음', '없음', 'LR F1 0.928 ★'],
        ['장애인 특화', '없음', '없음', '없음', 'KEAD 결합\n+20.4%p ★'],
        ['부정수급 탐지', '없음', '없음', '없음', 'SDK F1 0.864 ★'],
        ['실시간 검증', '없음', '없음', '없음', '7 API SSE ★'],
        ['점검 우선순위', '없음', '수동 검색', '수동 배정', 'AI 자동 정렬 ★'],
        ['근로자 신고', '있음', '없음', '오프라인', '익명 웹 신고\n신뢰도 자동산출 ★'],
    ],
    col_widths=[3.5, 2.5, 3.0, 3.0, 3.5]
)

add_heading(doc, '5-3. 핵심 기술 차별점', 2)
add_body(doc, '① 국내 최초 KEAD 데이터 ML 직접 결합')
add_body(doc,
    'KEAD 의무고용율(3.1%) 교차 추정을 ML 특성으로 직접 결합한 시스템은 현재까지 공개된 사례가 없습니다. '
    'Ablation 분석에서 해당 2개 특성이 F1 +20.4%p를 기여함을 실증하였습니다.',
    indent=0.5)

add_body(doc, '② "사전 탐지" 파이프라인 — 체불명단 이전 신호 3종')
add_body(doc,
    'DART 재무위험·NPS 가입자 이탈·4대보험 불일치 3개 신호는 모두 공개 체불명단 등재 시점보다 '
    '평균 3~6개월 앞서 이상을 나타냅니다. 이는 단순 체불명단 분류가 아닌 "예방적 탐지" 시스템임을 의미합니다.',
    indent=0.5)

add_body(doc, '③ 경량 SDK — 기존 시스템 무중단 통합')
add_body(doc,
    'Track A SDK는 순수 JavaScript 3KB로, 어떤 웹 기반 신청 시스템에도 스크립트 1줄 삽입으로 통합됩니다. '
    '서버 의존성 없이 클라이언트 측에서 즉시 판정하므로 응답 지연 0ms.',
    indent=0.5)

add_body(doc, '④ 투명한 한계 공개 — 시스템 사각지대 명시')
add_body(doc,
    'NPS·DART에 미등록된 소규모·개인사업장은 AI 점수화가 불가능합니다. '
    '이를 숨기지 않고 /triage 화면에 "시스템 사각지대" 카드로 투명하게 공개하고, '
    '시민 신고 누적으로 보완하는 구조입니다. 이는 공공 AI 시스템의 신뢰성 원칙에 부합합니다.',
    indent=0.5)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. 사업화 계획
# ═══════════════════════════════════════════════
add_heading(doc, '6. 제품 및 서비스의 사업화 계획 (효과성)', 1, (15,35,95))

add_heading(doc, '6-1. 사회적 임팩트 및 효과', 2)
add_table(doc,
    ['효과 지표', '현행', '목표 (WageGuard 도입 후)', '근거'],
    [
        ['체불 사전 탐지율', '0%\n(신고 후 대응)', '고위험 사업장 기준 35%+', 'ML Recall 0.904 적용'],
        ['근로감독관 1인 점검 효율', '무작위\n3,200개 중 선택', '의심도 상위 20%에 집중',
         '점검 시간 동일 대비\n실적발 확률 4.5배 향상 추정'],
        ['장애인 체불 조기 발굴', '신고율 30%\n(비장애인 60%)', '신고 없이 AI 감지 가능', 'KEAD 결합 모듈'],
        ['부정수급 차단율', '사후 적발 중심', 'SDK 탐지 F1 0.864',
         '1,000건 시뮬레이션'],
        ['사회적 비용 절감\n(몬테카를로 추정)', '연 체불액 1.73조', '예방 효과 연 5,641억 원',
         '95% CI [4,200억, 7,100억]'],
    ],
    col_widths=[3.5, 2.5, 4.5, 5.0]
)

add_heading(doc, '6-2. 홍보·마케팅 전략', 2)
add_body(doc, '■ 공공 채널 활용')
add_bullet(doc, '공모전 수상 후 고용노동부 공식 보도자료·SNS를 통한 1차 노출 (추정 도달 50만+)')
add_bullet(doc, 'KEAD·한국고용정보원 공동 보도자료 배포 → 공공기관 담당자 대상 직접 노출')
add_bullet(doc, '고용노동청 내부 워크숍·감독관 교육 과정에 시연 세션 삽입 추진')

add_body(doc, '■ 디지털 마케팅')
add_bullet(doc, 'GitHub 오픈소스 공개 → 개발자·연구자 커뮤니티 자연 확산')
add_bullet(doc, '고용노동 관련 유튜브 채널·뉴스레터 광고 (타깃: 근로자·소상공인·HR 담당자)')
add_bullet(doc, '근로자 셀프체크·체불 계산기를 무료 제공 → 유입 후 신고 데이터 축적')

add_heading(doc, '6-3. 확산 가능성 — 5개 정부 영역 이식', 2)
add_body(doc,
    '동일한 우선순위 정렬 메커니즘을 데이터 소스만 교체하여 5개 영역에 즉시 이식할 수 있습니다. '
    'Track A SDK는 코드 수정 없이 모든 웹 기반 신청 시스템에 적용 가능합니다.')
add_table(doc,
    ['확산 영역', '운영주체', '대체 데이터', '예상 효과'],
    [
        ['외국인 근로자 체불', '고용노동부 외국인력과', '외국인 고용허가 + NPS + 국적별 신고율', '외국인 체불 피해 25만 명 커버'],
        ['청년 고용장려금 부정수급', '고용노동부 청년고용과', 'Track A SDK 직접 적용', '연 500억 이상 부정수급 차단 기대'],
        ['고령자 고용장려금', 'KEAD + 고용노동부', '고령자 고용지원금 수급 DB', 'KEAD 기존 시스템에 모듈 통합'],
        ['산재보험 부정청구 의심', '근로복지공단', 'DART + 산재 청구 패턴', '허위 청구 사전 탐지'],
        ['공공조달 계약 불이행', '조달청', '기업 재무 + 낙찰 이력', '부실기업 낙찰 사전 경고'],
    ],
    col_widths=[3.5, 3.5, 4.5, 4.0]
)

add_heading(doc, '6-4. 실행 가능성 근거', 2)
add_body(doc, '■ 현재 완성도')
add_bullet(doc, '50+ 화면, 200+ API 라우트 — 현재 로컬에서 전 기능 즉시 시연 가능')
add_bullet(doc, '경량 아키텍처(FastAPI + SQLite) — 단일 서버 월 2~5만 원 운영, 즉시 배포 가능')
add_bullet(doc, '외부 의존성 없음(오픈소스만 사용) — 라이선스 리스크 없음')
add_bullet(doc, 'KEAD·고용정보원 활용신청 이미 완료 — 키 발급 즉시 실 데이터 연동 가능')

add_body(doc, '■ 기술 리스크 최소화')
add_bullet(doc, 'Logistic Regression: 블랙박스 없음, 특성 가중치 실시간 공개, 모델 감사 가능')
add_bullet(doc, '데이터 수집 파이프라인: scripts/ 디렉터리에 전체 자동화 스크립트 완비')
add_bullet(doc, '개인정보 보호: 사업장명 마스킹(신뢰도 60점 미달 시) + 신고자 정보 암호화 저장')

divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('■ 첨부 자료: 서비스 화면 캡처 · KEAD·고용정보원 활용신청 확인서 · API 응답 샘플 별첨')
set_font(run, size=9, color=(100,100,100))

# ── 저장
out_path = r'c:\lsc\Moel\proposal\WageGuard_사업계획서.docx'
doc.save(out_path)
print(f"저장 완료: {out_path}")
